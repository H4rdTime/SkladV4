from docx import Document

# Создаём новый документ
doc = Document()

# Заголовки
doc.add_paragraph("{{ org_name }}")
doc.add_paragraph("Коммерческое предложение № {{ estimate_number }} от {{ current_date_formatted }}")
doc.add_paragraph("{{ client_name }}")
doc.add_paragraph("{{ theme }}")

# Создаём таблицу 6 колонок: №, Наименование, Ед. изм., Кол-во, Цена, Сумма
table = doc.add_table(rows=2, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "№"
hdr_cells[1].text = "Наименование"
hdr_cells[2].text = "Ед. изм."
hdr_cells[3].text = "Кол-во"
hdr_cells[4].text = "Цена"
hdr_cells[5].text = "Сумма"

# Вторая строка — шаблон цикла
row_cells = table.rows[1].cells
row_cells[0].text = "{% for item in items %}\n{{ loop.index }}"
row_cells[1].text = "{{ item.product_name }}"
row_cells[2].text = "{{ item.unit }}"
row_cells[3].text = "{{ item.quantity }}"
row_cells[4].text = "{{ item.unit_price }}"
row_cells[5].text = "{{ item.total }}\n{% endfor %}"

# Итоговые поля
doc.add_paragraph("{{ total_sum_formatted }}")
doc.add_paragraph("{{ total_items_count }}")
doc.add_paragraph("{{ total_sum_in_words }}")
doc.add_paragraph("{{ valid_until_date_formatted }}")
doc.add_paragraph("{{ entrepreneur_name }}")

# Сохраняем файл
doc.save("commercial_proposal_template.docx")

print("Шаблон сохранён как commercial_proposal_template.docx")
